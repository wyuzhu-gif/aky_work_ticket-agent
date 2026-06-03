import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional
from uuid import uuid4

from common.logger import get_logger
from common.models import (
    RuleDocument, RuleDocumentSource, RuleDocumentStatus,
    DocumentRuleDocAssociation, ReviewRule, RiskLevel, RuleExample, RuleStatusEnum,
)
from database.rule_documents_repository import RuleDocumentsRepository
from database.rules_repository import RulesRepository
from database.db_client import SQLiteClient
from config.config import settings

logging = get_logger(__name__)

# ========== LLM Parsing Imports ==========
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field


class ParsedRule(BaseModel):
    name: str = Field(description="Rule name")
    description: str = Field(description="Rule description")
    risk_level: str = Field(description="Risk level: 高/中/低")
    examples: List[Dict[str, str]] = Field(default_factory=list, description="Examples: [{text, explanation}]")


class ParsedRulesOutput(BaseModel):
    rules: List[ParsedRule]


PARSE_SYSTEM_PROMPT = """你是一个规则文档解析专家。请从用户提供的文本中提取所有审核规则。
每条规则需要包含：
- name: 规则名称（简洁）
- description: 规则描述（详细说明检测什么问题）
- risk_level: 风险等级，只能是 高、中、低 之一
- examples: 示例列表，每个示例包含 text（问题文本）和 explanation（说明）

如果文本中无法提取出明确的规则，返回空列表。
严格按照 JSON 格式输出。
"""


class RuleDocumentsService:
    def __init__(self, rule_docs_repo: RuleDocumentsRepository, rules_repo: RulesRepository) -> None:
        self.rule_docs_repo = rule_docs_repo
        self.rules_repo = rules_repo

    # ========== CRUD ==========

    async def get_all_rule_documents(self) -> List[RuleDocument]:
        return await self.rule_docs_repo.get_all_rule_documents()

    async def get_rule_document(self, rule_doc_id: str) -> RuleDocument:
        return await self.rule_docs_repo.get_rule_document(rule_doc_id)

    async def get_rule_documents_by_ids(self, ids: List[str]) -> List[RuleDocument]:
        docs = []
        for doc_id in ids:
            try:
                doc = await self.rule_docs_repo.get_rule_document(doc_id)
                if doc.status == RuleDocumentStatus.active:
                    docs.append(doc)
            except ValueError:
                logging.warning(f"Rule document {doc_id} not found, skipping.")
        return docs

    async def delete_rule_document(self, rule_doc_id: str) -> None:
        doc = await self.rule_docs_repo.get_rule_document(rule_doc_id)
        # Delete associated file
        if doc.file_path:
            file_path = Path(doc.file_path)
            if file_path.exists():
                file_path.unlink()
        await self.rule_docs_repo.delete_rule_document(rule_doc_id)

    # ========== Upload ==========

    async def upload_rule_document(
        self,
        file_name: str,
        file_content: bytes,
        description: Optional[str] = None,
    ) -> RuleDocument:
        """Upload a rule document, extract text, and store it."""
        # Determine file type
        ext = Path(file_name).suffix.lower()
        if ext not in ('.pdf', '.docx', '.md', '.txt'):
            raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .md, .txt")

        file_type = ext.lstrip('.')

        # Save file
        docs_dir = Path(settings.rule_docs_dir)
        docs_dir.mkdir(parents=True, exist_ok=True)
        doc_id = str(uuid4())
        file_path = docs_dir / f"{doc_id}_{file_name}"
        file_path.write_bytes(file_content)
        logging.info(f"Rule document saved to {file_path}")

        # Extract text
        extracted_text = self._extract_text(file_path, file_type)
        logging.info(f"Extracted {len(extracted_text)} characters from {file_name}")

        # Create record
        rule_doc = RuleDocument(
            id=doc_id,
            name=file_name,
            description=description,
            file_path=str(file_path),
            file_type=file_type,
            source_type=RuleDocumentSource.context,
            extracted_text=extracted_text,
            status=RuleDocumentStatus.active,
            created_at=datetime.now(timezone.utc).isoformat(),
        )
        return await self.rule_docs_repo.create_rule_document(rule_doc)

    def _extract_text(self, file_path: Path, file_type: str) -> str:
        """Extract text from a file based on its type."""
        try:
            if file_type == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_type == 'docx':
                return self._extract_docx_text(file_path)
            elif file_type in ('md', 'txt'):
                return file_path.read_text(encoding='utf-8')
            else:
                return ""
        except Exception as e:
            logging.error(f"Failed to extract text from {file_path}: {e}")
            return ""

    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using PyMuPDF (fitz)."""
        import fitz
        text_parts = []
        doc = fitz.open(str(file_path))
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts)

    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document
        doc = Document(str(file_path))
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    # ========== LLM Parsing ==========

    async def parse_rule_document_with_llm(self, rule_doc_id: str) -> RuleDocument:
        """Parse a rule document into structured rules using LLM."""
        doc = await self.rule_docs_repo.get_rule_document(rule_doc_id)
        if not doc.extracted_text:
            raise ValueError(f"No extracted text for rule document {rule_doc_id}")

        # Call LLM to parse rules
        parsed_rules = await self._call_llm_parse(doc.extracted_text)
        logging.info(f"LLM parsed {len(parsed_rules)} rules from document {rule_doc_id}")

        # Create ReviewRule records
        created_rule_ids = []
        for pr in parsed_rules:
            risk_level = self._normalize_risk_level(pr.get("risk_level", "中"))
            examples = []
            for ex in pr.get("examples", []):
                examples.append({"text": ex.get("text", ""), "explanation": ex.get("explanation", "")})

            rule = ReviewRule(
                id=str(uuid4()),
                name=pr["name"],
                description=pr["description"],
                risk_level=risk_level,
                examples=[RuleExample(**ex) for ex in examples],
                status=RuleStatusEnum.active,
                created_at=datetime.now(timezone.utc).isoformat(),
            )
            created = await self.rules_repo.create_rule(rule)
            created_rule_ids.append(created.id)

        # Update rule document with parsed rule IDs
        fields = {
            "source_type": RuleDocumentSource.parsed,
            "parsed_rule_ids": created_rule_ids,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }
        return await self.rule_docs_repo.update_rule_document(rule_doc_id, fields)

    async def _call_llm_parse(self, text: str) -> List[Dict[str, Any]]:
        """Call LLM to parse text into structured rules."""
        from services.lc_pipeline import _init_llm_model
        llm = _init_llm_model()
        parser = PydanticOutputParser(pydantic_object=ParsedRulesOutput)

        # Truncate text if too long
        max_chars = 8000
        if len(text) > max_chars:
            text = text[:max_chars] + "\n... (文本已截断)"

        messages = [
            SystemMessage(content=PARSE_SYSTEM_PROMPT),
            HumanMessage(content=f"请从以下文本中提取审核规则：\n\n{text}\n\n{parser.get_format_instructions()}"),
        ]

        try:
            resp = await llm.ainvoke(messages)
            content = resp.content if hasattr(resp, "content") else resp
            if isinstance(content, list):
                content = "".join([c.get("text", "") if isinstance(c, dict) else str(c) for c in content])
            out = parser.parse(str(content))
            return [rule.model_dump() for rule in out.rules]
        except Exception as e:
            logging.error(f"LLM parsing failed: {e}")
            return []

    def _normalize_risk_level(self, level: str) -> RiskLevel:
        mapping = {"高": RiskLevel.high, "中": RiskLevel.medium, "低": RiskLevel.low}
        return mapping.get(level, RiskLevel.medium)

    # ========== Document Associations ==========

    async def get_document_rule_docs(self, doc_id: str) -> List[DocumentRuleDocAssociation]:
        return await self.rule_docs_repo.get_document_rule_docs(doc_id)

    async def get_enabled_rule_docs_for_document(self, doc_id: str) -> List[RuleDocument]:
        return await self.rule_docs_repo.get_enabled_rule_docs_for_document(doc_id)

    async def set_document_rule_doc(self, doc_id: str, rule_doc_id: str, enabled: bool) -> None:
        # Verify rule document exists
        await self.rule_docs_repo.get_rule_document(rule_doc_id)
        await self.rule_docs_repo.set_document_rule_doc(doc_id, rule_doc_id, enabled)
